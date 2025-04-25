import os
from dotenv import load_dotenv
import google.generativeai as genai
import json
from docx import Document
from datetime import datetime

load_dotenv()

def combined_output(result):
    """Combines all tool outputs and generates a structured research report, saving it as a .docx file."""

    # Load API key from environment
    api_key = os.getenv("GEMINI_API")
    genai.configure(api_key=api_key)

    # Prompt for Gemini model to structure the report
    prompt = """
    You are an expert research analyst with years of experience in compiling, analyzing, and presenting findings.
    Your task is to read the JSON data given to you which contains information from different tools:
    - Search engine data
    - Wikipedia knowledge
    - Current news highlights
    - Language model explanations

    Analyze all the data and synthesize it into a detailed, well-structured research report.
    
    Follow this structure:
    1. Title
    2. Abstract (Brief overview)
    3. Introduction
    4. Key Findings (from all tools)
    5. Analysis & Insights
    6. Conclusion
    7. References (if any URLs are present)

    Keep the tone academic and avoid repetition. Use bullet points and subheadings where appropriate.
    """

    try:
        # Format input data
        full_prompt = prompt + "\n\nHere is the data:\n" + json.dumps(result, indent=2)

        # Load Gemini LLM and get output
        llm = genai.GenerativeModel(model_name="gemini-2.5-flash-preview-04-17")
        output = llm.generate_content(full_prompt)
        report_text = output.text

        # Create directory if not exists
        output_dir = "reports"
        os.makedirs(output_dir, exist_ok=True)

        # Generate a timestamped filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"Research_Report_{timestamp}.docx"
        file_path = os.path.join(output_dir, filename)

        # Write the research report into a Word document
        document = Document()
        document.add_heading('Research Report', 0)

        for line in report_text.split("\n"):
            if line.strip().startswith("#"):
                document.add_heading(line.strip("#").strip(), level=1)
            else:
                document.add_paragraph(line.strip())

        document.save(file_path)
        print(f"report saved successfully at: {file_path}")
        return report_text

    except Exception as e:
        print("error generating or saving report:", e)
        return "An error occurred while generating the research report."
